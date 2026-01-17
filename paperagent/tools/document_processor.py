"""
PDF and document processing tools
"""

import os
from typing import Optional, List, Dict, Any
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
from loguru import logger


class PDFProcessor:
    """PDF document processor"""

    def extract_text(self, pdf_path: str) -> str:
        """
        Extract text from PDF

        Args:
            pdf_path: Path to PDF file

        Returns:
            Extracted text
        """
        try:
            text = []

            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)

            result = "\n\n".join(text)
            logger.info(f"Extracted {len(result)} characters from {pdf_path}")
            return result

        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return ""

    def extract_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract PDF metadata

        Args:
            pdf_path: Path to PDF file

        Returns:
            Dictionary of metadata
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', ''),
                }

                return metadata

        except Exception as e:
            logger.error(f"PDF metadata extraction error: {e}")
            return {}

    def extract_tables(self, pdf_path: str) -> List[List[List[str]]]:
        """
        Extract tables from PDF

        Args:
            pdf_path: Path to PDF file

        Returns:
            List of tables (each table is a list of rows)
        """
        try:
            tables = []

            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)

            logger.info(f"Extracted {len(tables)} tables from {pdf_path}")
            return tables

        except Exception as e:
            logger.error(f"PDF table extraction error: {e}")
            return []

    def extract_images(self, pdf_path: str, output_dir: str) -> List[str]:
        """
        Extract images from PDF

        Args:
            pdf_path: Path to PDF file
            output_dir: Directory to save extracted images

        Returns:
            List of paths to extracted images
        """
        try:
            os.makedirs(output_dir, exist_ok=True)
            image_paths = []

            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    if '/XObject' in page['/Resources']:
                        xobjects = page['/Resources']['/XObject'].get_object()

                        for obj_name in xobjects:
                            obj = xobjects[obj_name]

                            if obj['/Subtype'] == '/Image':
                                # Extract image data
                                size = (obj['/Width'], obj['/Height'])
                                data = obj.get_data()

                                # Save image
                                image_path = os.path.join(
                                    output_dir,
                                    f"page{page_num}_{obj_name[1:]}.png"
                                )

                                with open(image_path, 'wb') as img_file:
                                    img_file.write(data)

                                image_paths.append(image_path)

            logger.info(f"Extracted {len(image_paths)} images from {pdf_path}")
            return image_paths

        except Exception as e:
            logger.error(f"PDF image extraction error: {e}")
            return []

    def merge_pdfs(self, pdf_paths: List[str], output_path: str) -> bool:
        """
        Merge multiple PDFs into one

        Args:
            pdf_paths: List of PDF file paths
            output_path: Output PDF path

        Returns:
            Success boolean
        """
        try:
            pdf_merger = PyPDF2.PdfMerger()

            for pdf_path in pdf_paths:
                pdf_merger.append(pdf_path)

            with open(output_path, 'wb') as output_file:
                pdf_merger.write(output_file)

            logger.info(f"Merged {len(pdf_paths)} PDFs into {output_path}")
            return True

        except Exception as e:
            logger.error(f"PDF merge error: {e}")
            return False


class DOCXProcessor:
    """Microsoft Word document processor"""

    def create_document(
        self,
        title: str,
        sections: Dict[str, str],
        output_path: str
    ) -> bool:
        """
        Create a Word document

        Args:
            title: Document title
            sections: Dictionary of section_name: content
            output_path: Output .docx path

        Returns:
            Success boolean
        """
        try:
            doc = Document()

            # Add title
            doc.add_heading(title, level=0)

            # Add sections
            for section_name, content in sections.items():
                doc.add_heading(section_name, level=1)
                doc.add_paragraph(content)

            doc.save(output_path)
            logger.info(f"Word document created: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Word document creation error: {e}")
            return False

    def extract_text(self, docx_path: str) -> str:
        """
        Extract text from Word document

        Args:
            docx_path: Path to .docx file

        Returns:
            Extracted text
        """
        try:
            doc = Document(docx_path)

            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            result = "\n\n".join(text)
            logger.info(f"Extracted {len(result)} characters from {docx_path}")
            return result

        except Exception as e:
            logger.error(f"Word text extraction error: {e}")
            return ""

    def add_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]]
    ) -> None:
        """
        Add a table to Word document

        Args:
            doc: Document object
            headers: Table headers
            rows: Table rows
        """
        try:
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header

            # Rows
            for i, row in enumerate(rows, 1):
                row_cells = table.rows[i].cells
                for j, cell_value in enumerate(row):
                    row_cells[j].text = str(cell_value)

        except Exception as e:
            logger.error(f"Table addition error: {e}")


class TextProcessor:
    """Text processing utilities"""

    @staticmethod
    def split_into_sentences(text: str) -> List[str]:
        """
        Split text into sentences

        Args:
            text: Input text

        Returns:
            List of sentences
        """
        import re

        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]

        return sentences

    @staticmethod
    def split_into_paragraphs(text: str) -> List[str]:
        """
        Split text into paragraphs

        Args:
            text: Input text

        Returns:
            List of paragraphs
        """
        paragraphs = text.split('\n\n')
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        return paragraphs

    @staticmethod
    def count_words(text: str) -> int:
        """
        Count words in text

        Args:
            text: Input text

        Returns:
            Word count
        """
        words = text.split()
        return len(words)

    @staticmethod
    def extract_keywords(text: str, top_n: int = 10) -> List[str]:
        """
        Extract keywords from text using simple frequency analysis

        Args:
            text: Input text
            top_n: Number of top keywords to return

        Returns:
            List of keywords
        """
        from collections import Counter
        import re

        # Remove punctuation and convert to lowercase
        text = re.sub(r'[^\w\s]', '', text.lower())

        # Common stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be',
            'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
            'would', 'should', 'could', 'may', 'might', 'must', 'can', 'this',
            'that', 'these', 'those', 'it', 'its', 'they', 'them', 'their'
        }

        # Extract words
        words = text.split()
        words = [w for w in words if len(w) > 3 and w not in stop_words]

        # Count frequency
        word_freq = Counter(words)

        # Get top N keywords
        keywords = [word for word, _ in word_freq.most_common(top_n)]

        return keywords

    @staticmethod
    def calculate_readability(text: str) -> Dict[str, float]:
        """
        Calculate readability scores

        Args:
            text: Input text

        Returns:
            Dictionary of readability metrics
        """
        import re

        # Count sentences, words, and syllables (simplified)
        sentences = len(re.split(r'[.!?]+', text))
        words = len(text.split())

        # Simplified syllable count
        syllables = sum([TextProcessor._count_syllables(word) for word in text.split()])

        # Flesch Reading Ease
        if sentences > 0 and words > 0:
            flesch_reading_ease = 206.835 - 1.015 * (words / sentences) - 84.6 * (syllables / words)
        else:
            flesch_reading_ease = 0

        # Flesch-Kincaid Grade Level
        if sentences > 0 and words > 0:
            flesch_kincaid_grade = 0.39 * (words / sentences) + 11.8 * (syllables / words) - 15.59
        else:
            flesch_kincaid_grade = 0

        return {
            'flesch_reading_ease': max(0, min(100, flesch_reading_ease)),
            'flesch_kincaid_grade': max(0, flesch_kincaid_grade),
            'sentences': sentences,
            'words': words,
            'syllables': syllables
        }

    @staticmethod
    def _count_syllables(word: str) -> int:
        """
        Estimate syllable count in a word (simplified)

        Args:
            word: Input word

        Returns:
            Syllable count
        """
        word = word.lower()
        vowels = 'aeiouy'
        syllable_count = 0
        previous_was_vowel = False

        for char in word:
            is_vowel = char in vowels
            if is_vowel and not previous_was_vowel:
                syllable_count += 1
            previous_was_vowel = is_vowel

        # Adjust for silent 'e'
        if word.endswith('e'):
            syllable_count -= 1

        # Ensure at least one syllable
        return max(1, syllable_count)


# Export classes
__all__ = ["PDFProcessor", "DOCXProcessor", "TextProcessor"]
